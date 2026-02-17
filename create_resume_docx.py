from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the color of the hyperlink
    r.font.color.rgb = RGBColor(0, 0, 255)
    r.font.underline = True

    return hyperlink

def create_resume():
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Header ---
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run("NASHRAH ANAM FATHIMA")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = 'Times New Roman'

    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.add_run("📞 +91 6300126490 | ✉️ nashrahanam3@gmail.com | ")
    add_hyperlink(contact_paragraph, "LinkedIn", "https://linkedin.com/in/NashAnam")
    contact_paragraph.add_run(" | ")
    add_hyperlink(contact_paragraph, "GitHub", "https://github.com/NashAnam")
    contact_paragraph.runs[0].font.size = Pt(10)

    document.add_paragraph().add_run().add_break() # Spacer

    # --- Helper function for section headers ---
    def add_section_header(text):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.all_caps = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border


        # XML for bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- Education ---
    add_section_header("EDUCATION")
    
    table = document.add_table(rows=0, cols=2)
    table.autofit = True
    table.width = Inches(7.0) # Approx width
    
    # Row 1
    row_cells = table.add_row().cells
    p1 = row_cells[0].paragraphs[0]
    p1.add_run("B.Tech in AI & Data Science").bold = True
    p1.add_run("\nJNTU Hyderabad")
    
    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("2022 - 2026\nHyderabad, India")

    # Row 2
    row_cells = table.add_row().cells
    p1 = row_cells[0].paragraphs[0]
    p1.add_run("Intermediate (12th Grade)").bold = True
    p1.add_run("\nKakatiya Institution")
    
    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("2020 - 2022\n92.1%")

    # --- Technical Skills ---
    add_section_header("TECHNICAL SKILLS")
    
    skills = [
        ("Programming", "Python, Java, R, JavaScript"),
        ("Web Development", "HTML, CSS, React.js, Next.js, Streamlit, Flask"),
        ("AI/ML", "TensorFlow, PyTorch, Transformers, Scikit-learn, NumPy, Pandas"),
        ("Database", "SQLite, MongoDB, MySQL, PostgreSQL, Supabase"),
        ("Cloud & DevOps", "GCP, Firebase, Docker, Git, GitHub"),
        ("Tools", "VS Code, Jupyter, Postman, Android Studio, Power BI")
    ]

    for category, items in skills:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{category}: ")
        run.bold = True
        p.add_run(items)

    # --- Professional Experience ---
    add_section_header("PROFESSIONAL EXPERIENCE")

    # Job 1
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Full-Stack Developer Intern")
    r1.bold = True
    p.add_run(" | HealthOn (Telemedicine Platform)")
    
    # Date (using tabs for right align simplistically or just adding it)
    # Ideally a table or tab stops, but let's append it for simplicity or use a table row hidden
    # Using a 1x2 table for the header line ensures alignment
    
    # Re-doing Job 1 Header as Table
    p._element.getparent().remove(p._element) # Remove the paragraph we just made to replace with table
    
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    row_cells = table.rows[0].cells
    p1 = row_cells[0].paragraphs[0]
    r = p1.add_run("Full-Stack Developer Intern")
    r.bold = True
    p1.add_run(" | HealthOn (Telemedicine Platform)")
    
    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Jan 2026 - Present").italic = True

    # Job 1 Bullets
    bullets = [
        "Developing comprehensive telemedicine web application using Next.js, React, and Supabase for remote healthcare consultations",
        "Implemented real-time appointment scheduling system with automated reminders serving 500+ active users",
        "Built patient health tracking dashboard with vital signs monitoring and data visualization using Plotly",
        "Designed AI-powered progress reports analyzing patient health trends with personalized recommendations",
        "Established Row-Level Security (RLS) policies ensuring HIPAA-compliant data protection",
        "Optimized application performance achieving 40% faster page load times through code splitting"
    ]
    for b in bullets:
        p = document.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)


    # Job 2
    # Spacer
    document.add_paragraph().paragraph_format.space_after = Pt(6)

    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    row_cells = table.rows[0].cells
    p1 = row_cells[0].paragraphs[0]
    r = p1.add_run("Open Source Contributor")
    r.bold = True
    p1.add_run(" | AI Study Buddy")
    
    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Dec 2025 - Present").italic = True

    # Job 2 Bullets
    bullets = [
        "Built AI-powered study companion using Python, Streamlit, and Transformers for intelligent learning support",
        "Integrated DistilBART model for document summarization with TXT, PDF, and DOCX support",
        "Implemented keyword extraction using YAKE algorithm to identify key topics automatically",
        "Developed interactive flashcard system with difficulty ratings, shuffle mode, and progress tracking",
        "Applied bcrypt password hashing and SQL injection prevention for enterprise-grade security",
        "Achieved 5x faster summarization through model optimization and chunk-based processing"
    ]
    for b in bullets:
        p = document.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # --- Projects ---
    add_section_header("PROJECTS")

    def add_project(title, tech, date, bullets):
        table = document.add_table(rows=1, cols=2)
        table.autofit = True
        row_cells = table.rows[0].cells
        p1 = row_cells[0].paragraphs[0]
        r = p1.add_run(title)
        r.bold = True
        p1.add_run(f" | {tech}")
        
        p2 = row_cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(date).italic = True

        for b in bullets:
            p = document.add_paragraph(b, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        
        document.add_paragraph().paragraph_format.space_after = Pt(6)

    add_project("HealthOn - Telemedicine Platform", "Next.js, React, Supabase, Tailwind CSS", "Jan 2026 - Present", [
        "Comprehensive healthcare platform with patient-doctor matching, appointment scheduling, and telemedicine consultations",
        "Real-time vital signs tracking with trend analysis and AI-generated weekly progress reports",
        "Prescription management system with medication reminders and secure authentication with RLS"
    ])

    add_project("AI Study Buddy", "Python, Streamlit, Transformers, SQLite", "Dec 2025 - Present", [
        "AI-powered study tool with document summarization, keyword extraction, and customizable summary length",
        "Smart flashcard system with auto-generation, difficulty ratings, and spaced repetition",
        "Study session tracker with Pomodoro timer, analytics dashboard, and performance metrics"
    ])

    add_project("Attendance Tracker & Management System", "Python, OpenCV, ML", "Oct 2025 - Nov 2025", [
        "Built automated attendance system using facial recognition with 95% accuracy",
        "Reduced manual attendance time by 40% through automated CSV/PDF report generation"
    ])

    add_project("Intelligent Mental Health Companion", "Python, Flask, ML", "Sep 2025 - Oct 2025", [
        "Developed mental health support chatbot with sentiment analysis and mood tracking visualization",
        "Built web-based interface with Flask for seamless user interaction and data privacy"
    ])

    add_project("Smart Preventive Risk & Decision Support System", "Python, ML", "Aug 2025 - Present", [
        "Developing AI-driven health support system for early risk prediction and preventive care",
        "Integrating validated medical questionnaires (PHQ-9, GAD-7, FINDRISC) with ML models"
    ])

    # --- Achievements ---
    add_section_header("ACHIEVEMENTS & LEADERSHIP")
    achievements = [
        "Solved 50+ coding problems on LeetCode and HackerRank, strengthening Python & Java skills",
        "Designed and implemented production-ready AI applications focusing on education and mental health",
        "Active open-source contributor with focus on healthcare and education technology",
        "Engaged in hackathons and coding challenges to apply skills beyond academics",
        "GSoC 2026 aspirant building portfolio demonstrating AI/ML expertise and software engineering best practices"
    ]
    for a in achievements:
        p = document.add_paragraph(a, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # --- Personal Profile ---
    add_section_header("PERSONAL PROFILE")
    profile_table = document.add_table(rows=2, cols=2)
    profile_table.autofit = True
    
    # Needs to be a bit Manual to match grid
    # Row 1
    c = profile_table.rows[0].cells
    c[0].text = "Date of Birth: 27th April 2005"
    c[0].paragraphs[0].runs[0].bold = True # Attempt to bold label part... simpler to just overwrite
    c[0].paragraphs[0].clear()
    c[0].paragraphs[0].add_run("Date of Birth: ").bold = True
    c[0].paragraphs[0].add_run("27th April 2005")

    c[1].paragraphs[0].clear()
    c[1].paragraphs[0].add_run("Nationality: ").bold = True
    c[1].paragraphs[0].add_run("Indian")

    # Row 2
    c = profile_table.rows[1].cells
    c[0].paragraphs[0].clear()
    c[0].paragraphs[0].add_run("Languages: ").bold = True
    c[0].paragraphs[0].add_run("English, Urdu, Hindi")

    c[1].paragraphs[0].clear()
    c[1].paragraphs[0].add_run("Interests: ").bold = True
    c[1].paragraphs[0].add_run("AI/ML Research, Healthcare Tech")


    # --- Declaration ---
    document.add_paragraph().add_run().add_break()
    decl_p = document.add_paragraph()
    decl_p.add_run("Declaration: ").bold = True
    decl_p.add_run("I hereby declare that the above particulars are true and correct to the best of my knowledge.\n")
    decl_p.add_run("Place: ").bold = True
    decl_p.add_run("Hyderabad | ")
    decl_p.add_run("Date: ").bold = True
    decl_p.add_run("February 2026")
    decl_p.style.font.size = Pt(9)
    decl_p.style.font.italic = True

    # Save
    document.save('Nashrah_Anam_Resume.docx')
    print("Resume saved as Nashrah_Anam_Resume.docx")

if __name__ == "__main__":
    create_resume()
